from docx import Document
from docx.shared import Pt
from datetime import datetime
import os
 
class ReportGeneratorModel:
    def save_to_word(self,result_text):
            doc = Document()
            doc.add_heading('Результаты расчета конвейера', level=1)
            doc.add_paragraph(result_text)
                    
            # Сохраняем документ
            filename = f"Расчет конвеера{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
            
            # Проверяем папку для сохранения
            save_dir = "Отчеты"
            if not os.path.exists(save_dir):
                os.makedirs(save_dir)
            
            full_path = os.path.join(save_dir, filename)
            
            doc.save(full_path)